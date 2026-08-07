from pathlib import Path
import pandas as pd
from docx import Document

BASE_DIR = Path(__file__).resolve().parent
DOCX_PATH = BASE_DIR / "data" / "global_customer_inquiry_review.docx"
OUTPUT_DIR = BASE_DIR / "output"

PARAGRAPH_OUTPUT_PATH = OUTPUT_DIR / "word_paragraphs.csv"
TABLE_OUTPUT_PATH = OUTPUT_DIR / "word_table.csv"

def read_paragraphs(document: Document) -> pd.DataFrame:
    '''
    word 문서의 일반 문단 -> pd.DataFrame  변환
    '''
    paragraphs_records = []
    for paragraph_number, paragraph in enumerate(document.paragraphs):
        paragraph_text = paragraph.text.strip()

        if not paragraph_text:
            continue

        paragraphs_records.append(
            {
                'paragraph_number': paragraph_number,
                'style_name': paragraph.style.name,
                'paragraph_text': paragraph_text,
                'character_count': len(paragraph_text)
            }
        )
    return pd.DataFrame(paragraphs_records)


def read_tables(document: Document) -> pd.DataFrame:
    '''
    word 문서의 모든 표를 데이터 프레임으로 변환
    각 셀의 위치를 확인할 수 있도록 표 변호, 행 번호, 열번호 저장
    '''
    table_records = []
    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            for column_number, cell in enumerate(row.cells, start=1):
                cell_text = cell.text.strip()
                table_records.append(
                    {
                        'table_number': table_number,
                        'row_number': row_number,
                        'column_number': column_number,
                        'cell_text': cell_text
                    }
                )
    return pd.DataFrame(table_records)


def print_table_preview(table_df: pd.DataFrame) -> None:
    first_table = table_df[table_df['table_number']==1]
    first_table_matrix = first_table.pivot(
        index='row_number',
        columns='column_number',
        values='cell_text',
    )
    print('\n[첫 번째 표 미리보기]')
    print(first_table_matrix.to_string(index=False,header=False))


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(f'word 파일을 찾을 수 없습니다. : {DOCX_PATH}')

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    document = Document(DOCX_PATH)
    print(f'word file: {DOCX_PATH.name}')
    print(f'전체 일반 문단 수: {len(document.paragraphs)}')
    print(f'전체 표 수 : {len(document.tables)}')

    paragraphs_df = read_paragraphs(document)
    print('\n[문단 데이터 앞 부분')
    print(
        paragraphs_df[['paragraph_number', 'style_name', 'paragraph_text', 'character_count']].head(10).to_string(index=False)
    )

    table_df = read_tables(document)
    print_table_preview(table_df)

    paragraphs_df.to_csv(
        PARAGRAPH_OUTPUT_PATH,
        index=False,
        encoding='utf-8',
    )
    table_df.to_csv(
        TABLE_OUTPUT_PATH,
        index=False,
        encoding='utf-8',
    )

if __name__ == "__main__":
    main()